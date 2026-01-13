import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
from datetime import datetime
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
import json
import hashlib
import os
import copy
import re
import unicodedata
from collections import OrderedDict
import math

try:
    from docx import Document
    from docx.shared import Pt, Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from tkcalendar import Calendar
    TKCALENDAR_AVAILABLE = True
except ImportError:
    TKCALENDAR_AVAILABLE = False

class ModernButton(tk.Canvas):
    def __init__(self, parent, text, command, bg_color="#4A90E2", hover_color="#357ABD", 
                 text_color="white", width=150, height=40):
        super().__init__(parent, width=width, height=height, bg=parent.cget('bg'), 
                        highlightthickness=0)
        self.command = command
        self.bg_color = bg_color
        self.hover_color = hover_color
        self.text_color = text_color
        self.text = text
        
        self.rect = self.create_rectangle(0, 0, width, height, fill=bg_color, 
                                         outline="", tags="button")
        self.text_item = self.create_text(width//2, height//2, text=text, 
                                         fill=text_color, font=("Segoe UI", 11, "bold"),
                                         tags="button")
        
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
        self.bind("<Button-1>", self.on_click)
        self.tag_bind("button", "<Button-1>", self.on_click)
        
    def on_enter(self, e):
        self.itemconfig(self.rect, fill=self.hover_color)
        
    def on_leave(self, e):
        self.itemconfig(self.rect, fill=self.bg_color)
        
    def on_click(self, e):
        if self.command:
            self.command()

class EvaluatorTrainingApp:
    def __init__(self, root):
        self.root = root
        self.root.title("HTASO Umpire Evaluation Manager")
        self.root.geometry("1280x860")
        self.root.minsize(1100, 720)
        
        # Modern color palette
        self.colors = {
            'primary': '#1D3557',
            'primary_light': '#457B9D',
            'secondary': '#2A9D8F',
            'accent': '#E76F51',
            'bg': '#F1F5F9',
            'card': '#FFFFFF',
            'text': '#1F2937',
            'text_light': '#6B7280',
            'border': '#D1D5DB',
            'shadow': '#E0E7FF'
        }
        
        self.root.configure(bg=self.colors['bg'])
        self.root.grid_rowconfigure(1, weight=1)
        self.root.grid_rowconfigure(2, weight=0)
        self.root.grid_columnconfigure(0, weight=1)
        
        # Data directory
        self.data_dir = "evaluation_data"
        self.admin_file = os.path.join(self.data_dir, "admin_config.json")
        if not os.path.exists(self.data_dir):
            os.makedirs(self.data_dir)
        
        # Initialize admin password if not exists
        self.init_admin()
        
        # Storage for form data
        self.entries = {}
        self.ratings = {}
        self.rating_details = {}
        self.rating_scale = [
            ("5 - Excellent", 5),
            ("4 - Very Good", 4),
            ("3 - Satisfactory", 3),
            ("2 - Needs Improvement", 2),
            ("1 - Unsatisfactory", 1)
        ]
        self.optional_na_label = "Not Observed"
        self.rating_options = [label for label, _ in self.rating_scale]
        self.rating_value_map = {label: score for label, score in self.rating_scale}
        self.rating_value_map[self.optional_na_label] = None
        self.recommendation_indicators = []
        self.recommendation_default = "__none__"
        self.scroll_canvases = []
        self.root.bind_all("<MouseWheel>", self._handle_mousewheel, add="+")
        self.root.bind_all("<Button-4>", lambda e: self._handle_mousewheel(e, step=-1), add="+")
        self.root.bind_all("<Button-5>", lambda e: self._handle_mousewheel(e, step=1), add="+")
        self.root.bind("<Configure>", self._prune_scrollers, add="+")
        self.admin_window = None
        self.logo_path = Path(__file__).resolve().parent / "logo-150.png"
        
        # Check if admin mode
        self.is_admin = False
        
        # Load evaluation criteria from Excel template
        self.criteria_path = Path("Evaluator Training Eval form.xlsx")
        self.criteria_structure = self.load_eval_criteria_from_excel()
        
        # Create modern UI
        self.setup_styles()
        self.create_header()
        self.create_main_area()
        self.create_footer()
        
    def init_admin(self):
        """Initialize admin configuration"""
        if not os.path.exists(self.admin_file):
            # Default password: admin123
            default_hash = hashlib.sha256("admin123".encode()).hexdigest()
            with open(self.admin_file, 'w') as f:
                json.dump({"password_hash": default_hash}, f)
    
    def verify_admin_password(self, password):
        """Verify admin password"""
        if os.path.exists(self.admin_file):
            with open(self.admin_file, 'r') as f:
                config = json.load(f)
                password_hash = hashlib.sha256(password.encode()).hexdigest()
                return password_hash == config.get("password_hash")
        return False
    
    def setup_styles(self):
        """Configure ttk styles for a modern look"""
        self.style = ttk.Style()
        try:
            self.style.theme_use('clam')
        except tk.TclError:
            pass
        
        self.style.configure("Modern.TNotebook", background=self.colors['bg'], borderwidth=0)
        self.style.configure(
            "Modern.TNotebook.Tab",
            background=self.colors['card'],
            foreground=self.colors['text'],
            padding=(20, 12),
            font=("Segoe UI Semibold", 11),
            focuscolor=self.colors['card']
        )
        self.style.map(
            "Modern.TNotebook.Tab",
            background=[("selected", self.colors['primary'])],
            foreground=[("selected", "#FFFFFF")]
        )
        
        self.style.configure("Card.TFrame", background=self.colors['card'])
        self.style.configure("CardTitle.TLabel", background=self.colors['card'], foreground=self.colors['text'],
                             font=("Segoe UI Semibold", 13))
        self.style.configure("CardSubtitle.TLabel", background=self.colors['card'], foreground=self.colors['text_light'],
                             font=("Segoe UI", 10))
        
        self.style.configure(
            "Modern.TEntry",
            fieldbackground=self.colors['card'],
            foreground=self.colors['text'],
            bordercolor=self.colors['border'],
            lightcolor=self.colors['primary'],
            darkcolor=self.colors['border'],
            relief="flat",
            padding=6
        )
        self.style.map(
            "Modern.TEntry",
            fieldbackground=[("focus", "#FFFFFF")],
            bordercolor=[("focus", self.colors['primary'])],
            lightcolor=[("focus", self.colors['primary'])],
            darkcolor=[("focus", self.colors['primary'])]
        )
        
        self.style.configure(
            "Accent.TButton",
            background=self.colors['secondary'],
            foreground="#FFFFFF",
            font=("Segoe UI Semibold", 11),
            borderwidth=0,
            focusthickness=0,
            padding=(18, 10)
        )
        self.style.map(
            "Accent.TButton",
            background=[("active", self.colors['primary_light']), ("pressed", self.colors['primary'])],
            foreground=[("active", "#FFFFFF")]
        )
        
        self.style.configure(
            "Primary.TButton",
            background=self.colors['primary'],
            foreground="#FFFFFF",
            font=("Segoe UI Semibold", 11),
            borderwidth=0,
            focusthickness=0,
            padding=(18, 10)
        )
        self.style.map(
            "Primary.TButton",
            background=[("active", self.colors['primary_light']), ("pressed", self.colors['primary'])],
            foreground=[("active", "#FFFFFF")]
        )
        
        self.style.configure(
            "Neutral.TButton",
            background=self.colors['border'],
            foreground=self.colors['text'],
            font=("Segoe UI Semibold", 11),
            borderwidth=0,
            focusthickness=0,
            padding=(16, 10)
        )
        self.style.map(
            "Neutral.TButton",
            background=[("active", "#CBD5F5"), ("pressed", "#CBD5F5")],
            foreground=[("active", self.colors['text'])]
        )
    
    def create_main_area(self):
        """Create the central content area that hosts the tabs"""
        self.main_container = tk.Frame(self.root, bg=self.colors['bg'])
        self.main_container.grid(row=1, column=0, sticky="nsew", padx=24, pady=(12, 16))
        self.main_container.grid_rowconfigure(0, weight=1)
        self.main_container.grid_columnconfigure(0, weight=1)
        
        content_wrapper = tk.Frame(self.main_container, bg=self.colors['bg'])
        content_wrapper.grid(row=0, column=0, sticky="nsew")
        content_wrapper.grid_rowconfigure(0, weight=1)
        content_wrapper.grid_columnconfigure(0, weight=1)
        
        self.content_card = tk.Frame(
            content_wrapper,
            bg=self.colors['card'],
            bd=0,
            highlightbackground=self.colors['border'],
            highlightthickness=1
        )
        self.content_card.grid(row=0, column=0, sticky="nsew")
        self.create_tabs(self.content_card)
    
    def create_header(self):
        """Create modern header"""
        header = tk.Frame(self.root, bg=self.colors['primary'], height=100)
        header.grid(row=0, column=0, sticky="ew")
        header.grid_propagate(False)
        header.columnconfigure(0, weight=1)
        header.columnconfigure(1, weight=0)
        
        title_frame = tk.Frame(header, bg=self.colors['primary'])
        title_frame.grid(row=0, column=0, sticky="w", padx=30)
        
        if self.logo_path.exists():
            try:
                self.logo_image = tk.PhotoImage(file=str(self.logo_path))
                max_dimension = 100
                img_w = self.logo_image.width()
                if img_w > max_dimension:
                    scale = max(1, math.ceil(img_w / max_dimension))
                    self.logo_image = self.logo_image.subsample(scale, scale)
                img_h = self.logo_image.height()
                if img_h > max_dimension:
                    scale = max(1, math.ceil(img_h / max_dimension))
                    self.logo_image = self.logo_image.subsample(scale, scale)
                tk.Label(
                    title_frame,
                    image=self.logo_image,
                    bg=self.colors['primary']
                ).pack(side="left", padx=(0, 20), pady=(10, 0))
            except Exception:
                self.logo_image = None
        
        title_label = tk.Label(
            title_frame,
            text="HTASO Umpire Evaluation",
            font=("Segoe UI Semibold", 26),
            fg="#FFFFFF",
            bg=self.colors['primary']
        )
        title_label.pack(anchor="w", pady=(20, 0))
        
        subtitle_label = tk.Label(
            title_frame,
            text="Training, observations, and evaluations in one streamlined workspace.",
            font=("Segoe UI", 11),
            fg="#D1E3F8",
            bg=self.colors['primary']
        )
        subtitle_label.pack(anchor="w", pady=(6, 0))
        
        admin_frame = tk.Frame(header, bg=self.colors['primary'])
        admin_frame.grid(row=0, column=1, sticky="e", padx=30)
        
        tk.Label(
            admin_frame,
            text="Administrator",
            font=("Segoe UI", 9),
            fg="#C6DAF7",
            bg=self.colors['primary']
        ).pack(anchor="e", pady=(24, 6))
        
        admin_btn = ttk.Button(
            admin_frame,
            text="Manage Submissions",
            command=self.admin_login,
            style="Neutral.TButton"
        )
        admin_btn.pack(anchor="e")
    
    def create_tabs(self, parent):
        """Create modern tabbed interface"""
        tab_container = tk.Frame(parent, bg=self.colors['card'])
        tab_container.pack(fill="both", expand=True, padx=12, pady=12)
        
        self.notebook = ttk.Notebook(tab_container, style="Modern.TNotebook")
        self.notebook.pack(fill="both", expand=True)
        
        # Create tabs in specific order
        self.create_evaluation_tab()
        self.create_comments_tab()
        self.create_scale_criteria_tab()
        self.create_instructions_tab()

    def load_eval_criteria_from_excel(self):
        """Load evaluation criteria from the Excel template."""
        sections = []
        if not self.criteria_path.exists():
            return sections
        
        try:
            workbook = openpyxl.load_workbook(self.criteria_path, data_only=True)
        except Exception as exc:
            print(f"Failed to read evaluation template: {exc}")
            return sections
        
        sheet_name = "Eval. & Obser. Criteria"
        if sheet_name not in workbook.sheetnames:
            return sections
        
        sheet = workbook[sheet_name]
        current_section = None
        current_subsection = None
        last_item = None
        
        for row in sheet.iter_rows(min_row=1, max_col=11, values_only=True):
            col_a_raw = row[0]
            col_b_raw = row[1]
            col_c_raw = row[2]
            col_d_raw = row[3]
            col_e_raw = row[4]
            col_f_raw = row[5]
            others_raw = row[6:]
            
            col_a = self.clean_text(col_a_raw)
            col_b = self.clean_text(col_b_raw)
            col_c = self.clean_text(col_c_raw)
            col_d = self.clean_text(col_d_raw)
            col_e = self.clean_text(col_e_raw)
            col_f = self.clean_text(col_f_raw)
            other_text = [self.clean_text(val) for val in others_raw]
            
            # Detect high-level section (e.g. PLATE WORK)
            if col_e and not col_a and not col_b and not col_c:
                label = col_e.strip()
                lowered = label.lower()
                if not any(keyword in lowered for keyword in ("score", "out of", "pass or fail")):
                    current_section = {
                        "name": label.title(),
                        "raw_name": label,
                        "subsections": []
                    }
                    sections.append(current_section)
                    current_subsection = None
                    last_item = None
                continue
            
            if not current_section:
                continue
            
            # Detect subsection row (contains Score or Pass/Fail info)
            info_line = " ".join(filter(None, [col_c, col_d, col_e, col_f] + other_text)).lower()
            if col_a and (
                "score" in info_line or "pass or fail" in info_line
            ):
                numbers = []
                for value in [col_c_raw, col_d_raw, col_e_raw, col_f_raw] + list(others_raw):
                    if isinstance(value, (int, float)):
                        if value:
                            numbers.append(int(value))
                    elif isinstance(value, str):
                        match = re.search(r"(\d+)", value)
                        if match:
                            numbers.append(int(match.group(1)))
                max_score = max(numbers) if numbers else None
                current_subsection = {
                    "name": col_a,
                    "max_score": max_score,
                    "items": []
                }
                current_section["subsections"].append(current_subsection)
                last_item = None
                continue
            
            if not current_subsection:
                continue
            
            # Detect evaluation item rows
            if col_b:
                text = col_b
                col_a_clean = col_a.strip()
                is_numeric_label = col_a_clean.replace(".", "", 1).isdigit()
                is_new_item = True
                if last_item:
                    if not col_a_clean:
                        is_new_item = False
                    elif is_numeric_label:
                        first_char = text[:1]
                        if first_char and first_char.islower():
                            is_new_item = False
                        else:
                            lowered = text.lstrip().lower()
                            connector_prefixes = (
                                "and ", "or ", "nor ", "but ", "so ", "yet ",
                                "to ", "from ", "with ", "without ", "into ",
                                "onto ", "including ", "excluding "
                            )
                            if any(lowered.startswith(prefix) for prefix in connector_prefixes):
                                is_new_item = False
                    elif not col_a_clean:
                        is_new_item = False
                    else:
                        lowered = text.lstrip().lower()
                        if lowered and lowered[0].islower():
                            is_new_item = False
                if is_new_item:
                    key = self.slugify(
                        current_section["name"],
                        current_subsection["name"],
                        f"{len(current_subsection['items']) + 1:02d}"
                    )
                    allow_na = "n/a" in col_a.lower()
                    item = {
                        "key": key,
                        "text": text,
                        "allow_na": allow_na
                    }
                    current_subsection["items"].append(item)
                    last_item = item
                elif last_item:
                    joiner = "" if last_item["text"].endswith(("—", "-")) else " "
                    last_item["text"] = f"{last_item['text']}{joiner}{text}".strip()
                continue
        
        return self.merge_continued_sections(sections)
    
    def clean_text(self, value):
        """Normalize text extracted from Excel."""
        if value is None:
            return ""
        if isinstance(value, (int, float)):
            if float(value).is_integer():
                value = int(value)
            return str(value)
        text = str(value)
        text = unicodedata.normalize("NFKC", text)
        replacements = {
            "\u2013": "-",
            "\u2014": "-",
            "\u2018": "'",
            "\u2019": "'",
            "\u201c": '"',
            "\u201d": '"',
            "\u2022": "•",
            "�": "'"
        }
        for src, target in replacements.items():
            text = text.replace(src, target)
        text = text.replace("\n", " ").replace("\r", " ")
        text = re.sub(r"\s+", " ", text)
        return text.strip()
    
    def slugify(self, *parts):
        """Generate a consistent key for rating items."""
        joined = "_".join(filter(None, parts))
        normalized = unicodedata.normalize("NFKD", joined)
        ascii_text = normalized.encode("ascii", "ignore").decode()
        ascii_text = ascii_text.lower()
        ascii_text = re.sub(r"[^a-z0-9]+", "_", ascii_text)
        ascii_text = re.sub(r"_+", "_", ascii_text).strip("_")
        return ascii_text or "criterion"
    
    def merge_continued_sections(self, sections):
        """Combine sections that are continuations of the same group."""
        merged = {}
        order = []
        for section in sections:
            canonical = re.sub(r"\s*\(.*?\)\s*$", "", section["name"], flags=re.IGNORECASE).strip()
            if canonical not in merged:
                merged[canonical] = {
                    "name": canonical,
                    "raw_name": section["raw_name"],
                    "subsections": []
                }
                order.append(canonical)
            else:
                existing_raw = merged[canonical]["raw_name"]
                if section["raw_name"] not in existing_raw:
                    merged[canonical]["raw_name"] = f"{existing_raw}, {section['raw_name']}"
            merged[canonical]["subsections"].extend(section["subsections"])
        return [merged[name] for name in order]
    
    def compute_section_summary(self, ratings):
        """Calculate per-section and overall scores from rating entries."""
        totals = OrderedDict()
        total_score = 0
        total_possible = 0
        
        for item in ratings:
            section = item.get("section") or "General"
            score = item.get("score")
            if isinstance(score, str):
                try:
                    score = float(score)
                except ValueError:
                    score = None
            if not isinstance(score, (int, float)):
                score = None
            if section not in totals:
                totals[section] = {"score": 0.0, "count": 0}
            if score is not None and score > 0:
                totals[section]["score"] += score
                totals[section]["count"] += 1
                total_score += score
                total_possible += 5
        
        section_list = []
        for section, data in totals.items():
            possible = data["count"] * 5
            percentage = (data["score"] / possible) if possible else 0
            section_list.append({
                "section": section,
                "score": data["score"],
                "possible": possible,
                "percentage": percentage
            })
        
        return section_list, total_score, total_possible
    
    def create_evaluation_tab(self):
        """Create the evaluation tab with criteria sourced from Excel"""
        tab = tk.Frame(self.notebook, bg=self.colors['bg'])
        self.notebook.add(tab, text="Evaluation")
        tab.grid_rowconfigure(0, weight=1)
        tab.grid_columnconfigure(0, weight=1)
        
        canvas = tk.Canvas(tab, bg=self.colors['bg'], highlightthickness=0)
        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        self.eval_frame = tk.Frame(canvas, bg=self.colors['bg'])
        canvas.create_window((0, 0), window=self.eval_frame, anchor="nw")
        
        self.eval_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        self.enable_mousewheel(canvas)
        
        # Basic information card
        info_card = self.create_card(
            self.eval_frame,
            "Evaluator & Session Details",
            subtitle="Complete these required fields before scoring."
        )
        info_card.pack(fill="x", padx=30, pady=20)
        info_body = info_card.body
        
        info_grid = tk.Frame(info_body, bg=self.colors['card'])
        info_grid.pack(fill="x", pady=(8, 0))
        for column in range(3):
            info_grid.columnconfigure(column, weight=1)
        
        info_fields = [
            ("Evaluator Name *", "evaluator_name"),
            ("Trainer Name *", "trainer_name"),
            ("Training Date *", "training_date"),
            ("Observation Date", "observation_date"),
            ("Training Location", "training_location"),
            ("Type of Evaluation", "eval_type")
        ]
        
        for index, (label, key) in enumerate(info_fields):
            row = index // 3
            column = index % 3
            field_frame = tk.Frame(info_grid, bg=self.colors['card'])
            field_frame.grid(row=row, column=column, sticky="ew", padx=10, pady=10)
            
            tk.Label(
                field_frame,
                text=label,
                font=("Segoe UI", 10, "bold"),
                fg=self.colors['text'],
                bg=self.colors['card']
            ).pack(anchor="w")
            
            if key == "training_date":
                container = tk.Frame(field_frame, bg=self.colors['card'])
                container.pack(fill="x", pady=(6, 0))
                entry = ttk.Entry(container, font=("Segoe UI", 11), style="Modern.TEntry")
                entry.pack(side="left", fill="x", expand=True)
                if TKCALENDAR_AVAILABLE:
                    tk.Button(
                        container,
                        text="📅",
                        command=lambda e=entry: self.open_date_picker(e),
                        font=("Segoe UI", 10),
                        bg="white",
                        fg=self.colors['primary'],
                        relief="flat",
                        padx=8,
                        pady=4,
                        cursor="hand2"
                    ).pack(side="left", padx=(6, 0))
                self.entries[key] = entry
                if not TKCALENDAR_AVAILABLE:
                    tk.Label(
                        field_frame,
                        text="Install tkcalendar for date picker",
                        font=("Segoe UI", 8, "italic"),
                        bg=self.colors['card'],
                        fg=self.colors['text_light']
                    ).pack(anchor="w", pady=(4, 0))
            else:
                entry = ttk.Entry(field_frame, font=("Segoe UI", 11), style="Modern.TEntry")
                entry.pack(fill="x", pady=(6, 0))
                self.entries[key] = entry
        
        # Recommendation card
        rec_card = self.create_card(
            self.eval_frame,
            "Overall Recommendation",
            subtitle="Select the option that best reflects the evaluator's readiness."
        )
        rec_card.pack(fill="x", padx=30, pady=(0, 20))
        rec_body = rec_card.body
        self.recommendation = tk.StringVar(value=self.recommendation_default)
        self.recommendation_indicators = []
        recommendations = [
            ("Approved for Independent Evaluation", self.colors['secondary']),
            ("Approved with Additional Training Required", "#F4A261"),
            ("Requires Further Training Before Approval", "#F97316"),
            ("Not Approved - Significant Concerns", self.colors['accent'])
        ]
        
        for text, color in recommendations:
            row = tk.Frame(rec_body, bg=self.colors['card'])
            row.pack(fill="x", pady=6)
            indicator = tk.Frame(row, bg=self.colors['border'], width=6, height=32)
            indicator.pack(side="left", padx=(0, 12))
            indicator.pack_propagate(False)
            
            rb = tk.Radiobutton(
                row,
                text=text,
                variable=self.recommendation,
                value=text,
                font=("Segoe UI", 11),
                bg=self.colors['card'],
                fg=self.colors['text'],
                activebackground=self.colors['card'],
                selectcolor=self.colors['card'],
                cursor="hand2",
                command=self.update_recommendation_indicators
            )
            rb.pack(side="left", fill="x", expand=True)
            rb.deselect()
            self.recommendation_indicators.append((indicator, color, text))
        
        self.update_recommendation_indicators()
        
        # Evaluation criteria sourced from Excel
        if not self.criteria_structure:
            fallback = self.create_card(
                self.eval_frame,
                "Evaluation Criteria Unavailable",
                subtitle="The Excel template could not be read. Please ensure it is closed and accessible."
            )
            fallback.pack(fill="x", padx=30, pady=10)
            tk.Label(
                fallback.body,
                text="Close the Excel workbook and reopen this application, or contact an administrator for assistance.",
                font=("Segoe UI", 11),
                fg=self.colors['text_light'],
                bg=self.colors['card'],
                wraplength=900,
                justify="left"
            ).pack(fill="x", pady=12)
            return
        
        for section in self.criteria_structure:
            section_card = self.create_card(
                self.eval_frame,
                section["name"],
                subtitle=f"Sourced from template section: {section['raw_name']}"
            )
            section_card.pack(fill="x", padx=30, pady=12)
            section_body = section_card.body
            
            for subsection in section["subsections"]:
                subsection_frame = tk.Frame(section_body, bg=self.colors['card'])
                subsection_frame.pack(fill="x", pady=(10, 18))
                
                header = tk.Frame(subsection_frame, bg=self.colors['card'])
                header.pack(fill="x")
                
                tk.Label(
                    header,
                    text=subsection["name"],
                    font=("Segoe UI Semibold", 12),
                    fg=self.colors['text'],
                    bg=self.colors['card']
                ).pack(side="left")
                
                if subsection.get("max_score"):
                    tk.Label(
                        header,
                        text=f"Target: {subsection['max_score']} points",
                        font=("Segoe UI", 10),
                        fg=self.colors['text_light'],
                        bg=self.colors['card']
                    ).pack(side="right")
                
                for item in subsection["items"]:
                    row_frame = tk.Frame(subsection_frame, bg=self.colors['card'])
                    row_frame.pack(fill="x", pady=6)
                    
                    text_label = tk.Label(
                        row_frame,
                        text=item["text"],
                        font=("Segoe UI", 10),
                        fg=self.colors['text'],
                        bg=self.colors['card'],
                        justify="left",
                        wraplength=900
                    )
                    text_label.pack(side="left", fill="x", expand=True, padx=(0, 12))
                    
                    options = list(self.rating_options)
                    if item.get("allow_na") and self.optional_na_label not in options:
                        options = options + [self.optional_na_label]
                    display_options = ["Select result"] + options
                    
                    var = tk.StringVar(value=display_options[0])
                    combo = ttk.Combobox(
                        row_frame,
                        textvariable=var,
                        state="readonly",
                        values=display_options,
                        width=24
                    )
                    combo.pack(side="right")
                    combo.current(0)
                    
                    self.ratings[item["key"]] = var
                    self.rating_details[item["key"]] = {
                        "section": section["name"],
                        "subsection": subsection["name"],
                        "prompt": item["text"]
                    }
    
    def create_comments_tab(self):
        """Create modern comments tab"""
        tab = tk.Frame(self.notebook, bg=self.colors['bg'])
        self.notebook.add(tab, text="💬 Comments")
        
        canvas = tk.Canvas(tab, bg=self.colors['bg'], highlightthickness=0)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
        scrollbar.pack(side="right", fill="y")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        frame = tk.Frame(canvas, bg=self.colors['bg'])
        canvas.create_window((0, 0), window=frame, anchor="nw")
        frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        self.enable_mousewheel(canvas)
        
        content = tk.Frame(frame, bg=self.colors['bg'])
        content.pack(fill="both", expand=True, padx=30, pady=30)
        
        sections = [
            ("💪 Strengths Observed", "strengths_text", 
             "Document specific examples of strong performance and positive attributes"),
            ("📈 Areas for Improvement", "improvement_text",
             "Identify specific areas needing development with constructive feedback"),
            ("🎯 Recommendations for Development", "development_text",
             "Suggest specific actions, training, or resources for growth"),
            ("📝 Overall Assessment Comments", "overall_text",
             "Provide additional observations, context, or summary comments")
        ]
        
        for title, attr_name, placeholder in sections:
            card = self.create_card(content, title)
            card.pack(fill="both", expand=True, pady=10)
            card_content = card.body
            
            tk.Label(
                card_content,
                text=placeholder,
                font=("Segoe UI", 9, "italic"),
                fg=self.colors['text_light'],
                bg=self.colors['card']
            ).pack(anchor="w", pady=(0, 6))
            
            text_widget = scrolledtext.ScrolledText(
                card_content,
                height=6,
                font=("Segoe UI", 10),
                relief="flat",
                wrap="word"
            )
            text_widget.pack(fill="both", expand=True)
            text_widget.configure(
                background=self.colors['card'],
                foreground=self.colors['text'],
                insertbackground=self.colors['text'],
                highlightthickness=1,
                highlightbackground=self.colors['border'],
                highlightcolor=self.colors['primary']
            )
            setattr(self, attr_name, text_widget)
    
    def create_scale_criteria_tab(self):
        """Create modern scale criteria tab"""
        tab = tk.Frame(self.notebook, bg=self.colors['bg'])
        self.notebook.add(tab, text="📏 Scale Criteria")
        
        canvas = tk.Canvas(tab, bg=self.colors['bg'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
        frame = tk.Frame(canvas, bg=self.colors['bg'])
        
        frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        self.enable_mousewheel(canvas)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        title_card = self.create_card(frame, "📊 Rating Scale Criteria")
        title_card.pack(fill="x", padx=30, pady=20)
        tk.Label(
            title_card.body,
            text="Use this scale to keep scoring consistent across every observation.",
            font=("Segoe UI", 11),
            fg=self.colors['text_light'],
            bg=self.colors['card'],
            wraplength=860,
            justify="left"
        ).pack(fill="x")
        
        scales = [
            ("5 - EXCELLENT", "#2ECC71", 
             "• Consistently exceeds expectations\n• Demonstrates exceptional understanding and skill\n• Requires no additional support or guidance\n• Serves as a model for other evaluators\n• Shows mastery of all evaluation aspects"),
            
            ("4 - VERY GOOD", "#3498DB",
             "• Consistently meets and often exceeds expectations\n• Demonstrates strong understanding and competence\n• Requires minimal guidance\n• Performs reliably with high quality results\n• Shows proficiency in most evaluation aspects"),
            
            ("3 - SATISFACTORY", "#F39C12",
             "• Meets basic expectations\n• Demonstrates adequate understanding and capability\n• May require occasional guidance or support\n• Performs competently with acceptable results\n• Shows fundamental competence in evaluation tasks"),
            
            ("2 - NEEDS IMPROVEMENT", "#E67E22",
             "• Partially meets expectations\n• Shows gaps in understanding or application\n• Requires regular guidance and support\n• Performance is inconsistent or below standard\n• Needs additional training in several areas"),
            
            ("1 - UNSATISFACTORY", "#E74C3C",
             "• Does not meet minimum expectations\n• Demonstrates significant gaps in knowledge or skill\n• Requires extensive support and retraining\n• Performance is consistently below acceptable standards\n• Not ready for independent evaluation work")
        ]
        
        for rating, color, description in scales:
            scale_card = self.create_card(frame, rating, accent_color=color)
            scale_card.pack(fill="x", padx=30, pady=12)
            tk.Label(
                scale_card.body,
                text=description,
                font=("Segoe UI", 10),
                fg=self.colors['text'],
                bg=self.colors['card'],
                justify="left"
            ).pack(anchor="w", pady=(0, 6))
    
    def create_instructions_tab(self):
        """Create modern instructions tab"""
        tab = tk.Frame(self.notebook, bg=self.colors['bg'])
        self.notebook.add(tab, text="📖 Instructions")
        
        canvas = tk.Canvas(tab, bg=self.colors['bg'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(tab, orient="vertical", command=canvas.yview)
        frame = tk.Frame(canvas, bg=self.colors['bg'])
        
        frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        self.enable_mousewheel(canvas)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        instruction_sections = [
            ("📋 PURPOSE", "This form is designed to assess the performance and readiness of evaluators during or after their training period. It helps ensure that evaluators meet the required standards before conducting independent evaluations."),
            
            ("📝 HOW TO USE THIS FORM", 
             "1. Complete Basic Information - Fill in the evaluator's name, training date, trainer name, and location\n\n2. Review the Scale Criteria - Familiarize yourself with the 5-point rating scale\n\n3. Complete the Evaluation - Rate each criterion based on observed performance\n\n4. Provide Detailed Comments - Document specific strengths and areas for improvement\n\n5. Make a Final Recommendation - Select the appropriate recommendation\n\n6. Submit the Evaluation - Click 'Submit Evaluation' to save securely"),
            
            ("⚠️ IMPORTANT NOTES",
             "• Be thorough and specific in your observations\n• Support ratings with concrete examples\n• Use the full range of the rating scale\n• Maintain confidentiality of evaluation results\n• Review completed form with the evaluator when appropriate")
        ]
        
        for title, content in instruction_sections:
            card = self.create_card(frame, title)
            card.pack(fill="x", padx=30, pady=15)
            
            text = tk.Label(card.body, text=content, font=("Segoe UI", 10),
                          justify="left", anchor="w", bg=self.colors['card'],
                          fg=self.colors['text'], wraplength=900)
            text.pack(anchor="w", pady=15)
    
    def create_card(self, parent, title, subtitle=None, accent_color=None):
        """Create a modern card UI element"""
        card = tk.Frame(
            parent,
            bg=self.colors['card'],
            highlightbackground=self.colors['border'],
            highlightthickness=1,
            bd=0
        )
        
        accent = tk.Frame(card, bg=accent_color or self.colors['primary'], height=3)
        accent.pack(fill="x", side="top")
        
        header = tk.Frame(card, bg=self.colors['card'])
        header.pack(fill="x", padx=24, pady=(18, 8))
        
        tk.Label(
            header,
            text=title,
            font=("Segoe UI Semibold", 14),
            fg=self.colors['text'],
            bg=self.colors['card']
        ).pack(anchor="w")
        
        if subtitle:
            tk.Label(
                header,
                text=subtitle,
                font=("Segoe UI", 10),
                fg=self.colors['text_light'],
                bg=self.colors['card']
            ).pack(anchor="w", pady=(4, 0))
        
        separator = tk.Frame(card, bg=self.colors['border'], height=1)
        separator.pack(fill="x", padx=24, pady=(0, 12))
        
        body = tk.Frame(card, bg=self.colors['card'])
        body.pack(fill="both", expand=True, padx=24, pady=(0, 20))
        card.body = body
        
        return card
    
    def create_footer(self):
        """Create modern footer with action buttons"""
        footer = tk.Frame(self.root, bg=self.colors['bg'], height=90)
        footer.grid(row=2, column=0, sticky="ew", padx=24, pady=(0, 20))
        footer.grid_propagate(False)
        
        footer.columnconfigure(0, weight=1)
        
        button_frame = tk.Frame(
            footer,
            bg=self.colors['card'],
            highlightbackground=self.colors['border'],
            highlightthickness=1
        )
        button_frame.place(relx=0.5, rely=0.5, anchor="center")
        
        ModernButton(
            button_frame,
            "Submit Evaluation",
            self.submit_evaluation,
            bg_color=self.colors['secondary'],
            hover_color="#45B56A",
            width=190,
            height=46
        ).pack(side="left", padx=12, pady=12)
        
        ModernButton(
            button_frame,
            "Export Word Report",
            self.export_report_word,
            bg_color=self.colors['primary'],
            hover_color=self.colors['primary_light'],
            width=190,
            height=46
        ).pack(side="left", padx=12, pady=12)
        
        ModernButton(
            button_frame,
            "Export PDF Report",
            self.export_report_pdf,
            bg_color=self.colors['accent'],
            hover_color="#d75a3b",
            width=190,
            height=46
        ).pack(side="left", padx=12, pady=12)
        
        ModernButton(
            button_frame,
            "Clear Form",
            self.clear_form,
            bg_color="#94A3B8",
            hover_color="#64748B",
            width=160,
            height=46
        ).pack(side="left", padx=12, pady=12)
    
    def enable_mousewheel(self, canvas):
        """Register a canvas to respond to global mouse wheel events."""
        if canvas not in self.scroll_canvases:
            self.scroll_canvases.append(canvas)
            canvas.bind("<Destroy>", lambda e, c=canvas: self._remove_scroller(c), add="+")
    
    def _remove_scroller(self, canvas):
        if canvas in self.scroll_canvases:
            self.scroll_canvases.remove(canvas)
    
    def _handle_mousewheel(self, event, step=None):
        """Dispatch mouse wheel events to the canvas underneath the cursor."""
        self._prune_scrollers()
        
        target = None
        widget = getattr(event, "widget", None)
        while widget is not None:
            if widget in self.scroll_canvases:
                target = widget
                break
            widget = getattr(widget, "master", None)
        
        pointer_x = getattr(event, "x_root", None)
        pointer_y = getattr(event, "y_root", None)
        if pointer_x is None or pointer_y is None:
            pointer_x = self.root.winfo_pointerx()
            pointer_y = self.root.winfo_pointery()
        if target is None and pointer_x is not None and pointer_y is not None:
            for canvas in reversed(self.scroll_canvases):
                if not canvas.winfo_exists():
                    continue
                x0 = canvas.winfo_rootx()
                y0 = canvas.winfo_rooty()
                x1 = x0 + canvas.winfo_width()
                y1 = y0 + canvas.winfo_height()
                if x0 <= pointer_x < x1 and y0 <= pointer_y < y1:
                    target = canvas
                    break
        
        if target is None:
            return
        
        if step is None:
            if hasattr(event, "delta") and event.delta:
                delta = int(-1 * (event.delta / 120))
                if delta == 0:
                    delta = -1 if event.delta > 0 else 1
            else:
                delta = -1
        else:
            delta = step
        
        target.yview_scroll(delta, "units")
    
    def _prune_scrollers(self, *_):
        self.scroll_canvases = [c for c in self.scroll_canvases if c.winfo_exists()]
    
    def update_recommendation_indicators(self):
        """Refresh recommendation accent bars to reflect current selection."""
        current = self.recommendation.get()
        for indicator, color, value in self.recommendation_indicators:
            indicator.configure(bg=color if value == current else self.colors['border'])
    
    def collect_ratings(self):
        """Gather current rating selections with summary statistics."""
        rating_entries = []
        score_counts = {label: 0 for label in self.rating_value_map.keys()}
        rated_items = 0
        
        for key, var in self.ratings.items():
            selection = var.get()
            if selection == "Select result":
                selection = ""
            detail = self.rating_details.get(key, {})
            score_value = self.rating_value_map.get(selection)
            
            if selection and selection in score_counts:
                score_counts[selection] += 1
            if score_value is not None:
                rated_items += 1
            
            rating_entries.append({
                "key": key,
                "section": detail.get("section"),
                "subsection": detail.get("subsection"),
                "prompt": detail.get("prompt"),
                "selection": selection or None,
                "score": score_value
            })
        
        section_totals, total_score, total_possible = self.compute_section_summary(rating_entries)
        average_score = (total_score / total_possible) if total_possible else 0.0
        summary_section_totals = [
            {
                "section": entry["section"],
                "score": entry["score"],
                "possible": entry["possible"],
                "percentage": entry["percentage"]
            }
            for entry in section_totals
        ]

        return {
            "entries": rating_entries,
            "average": average_score,
            "rated_count": rated_items,
            "score_counts": score_counts,
            "section_totals": summary_section_totals,
            "total_score": total_score,
            "total_possible": total_possible
        }
    
    def build_form_evaluation_data(self):
        """Collect current form data and computed summaries without persisting."""
        summary = self.collect_ratings()
        recommendation_value = self.recommendation.get()
        if recommendation_value == self.recommendation_default:
            recommendation_value = ""
        eval_data = {
            "evaluator_name": self.entries["evaluator_name"].get(),
            "training_date": self.entries["training_date"].get(),
            "trainer_name": self.entries["trainer_name"].get(),
            "training_location": self.entries["training_location"].get(),
            "observation_date": self.entries["observation_date"].get(),
            "eval_type": self.entries["eval_type"].get(),
            "ratings": summary["entries"],
            "recommendation": recommendation_value,
            "average_score": summary["average"],
            "score_percentage": summary["average"] * 100,
            "rated_item_count": summary["rated_count"],
            "score_counts": summary["score_counts"],
            "section_totals": summary["section_totals"],
            "total_score": summary["total_score"],
            "total_possible": summary["total_possible"],
            "comments": {
                "strengths": self.strengths_text.get("1.0", "end-1c"),
                "improvement": self.improvement_text.get("1.0", "end-1c"),
                "development": self.development_text.get("1.0", "end-1c"),
                "overall": self.overall_text.get("1.0", "end-1c")
            }
        }
        return eval_data, summary
    
    def prepare_report_context(self, data, prepared_on=None):
        """Normalize evaluation data and prepare context for reporting."""
        record = self.normalize_evaluation_data(data)
        section_summary = record.get("section_totals") or []
        total_score = record.get("total_score", 0)
        total_possible = record.get("total_possible", 0)
        if not section_summary or total_possible is None:
            section_summary, total_score, total_possible = self.compute_section_summary(record.get("ratings", []))
        average_score = (total_score / total_possible) if total_possible else 0
        average_display = f"{average_score:.0%}" if total_possible else "N/A"
        overall_raw = f"{total_score:.0f}/{total_possible}" if total_possible else "0/0"
        
        grouped = OrderedDict()
        for item in record.get("ratings", []):
            section_name = item.get("section") or "General"
            subsection_name = item.get("subsection") or "Criteria"
            grouped.setdefault(section_name, OrderedDict()).setdefault(subsection_name, []).append(item)
        
        comments_source = record.get("comments", {})
        def fetch_comment(*keys):
            for key in keys:
                if key in comments_source and comments_source[key]:
                    return comments_source[key]
            return ""
        
        context = {
            "prepared_on": prepared_on or datetime.now(),
            "evaluator_name": record.get("evaluator_name", "Unknown"),
            "trainer_name": record.get("trainer_name", "Unknown"),
            "training_date": record.get("training_date", "N/A"),
            "training_location": record.get("training_location", "N/A"),
            "submission_date": self.format_datetime_display(record.get("submission_date", "N/A")),
            "observation_date": record.get("observation_date", ""),
            "eval_type": record.get("eval_type", ""),
            "rated_items": record.get("rated_item_count", 0),
            "average_display": average_display,
            "overall_raw": overall_raw,
            "average_score": average_score,
            "section_totals": section_summary,
            "total_score": total_score,
            "total_possible": total_possible,
            "grouped_ratings": grouped,
            "recommendation": record.get("recommendation") or "Not Selected",
            "comments": {
                "Strengths": fetch_comment("strengths", "STRENGTHS:", "Strengths"),
                "Areas for Improvement": fetch_comment("improvement", "IMPROVEMENTS:", "Areas for Improvement"),
                "Development Recommendations": fetch_comment("development", "DEVELOPMENT:", "Development Recommendations"),
                "Overall Comments": fetch_comment("overall", "OVERALL:", "Overall")
            }
        }
        return context
    
    def normalize_evaluation_data(self, data):
        """Ensure legacy evaluation records conform to the current data structure."""
        normalized = copy.deepcopy(data)
        
        ratings = normalized.get("ratings", [])
        if isinstance(ratings, dict):
            converted = []
            for key, value in ratings.items():
                converted.append({
                    "key": key,
                    "section": "Legacy",
                    "subsection": "Criteria",
                    "prompt": key.replace('_', ' ').title(),
                    "selection": f"{value}/5" if isinstance(value, (int, float)) else value,
                    "score": value if isinstance(value, (int, float)) else None
                })
            normalized["ratings"] = converted
        elif not isinstance(ratings, list):
            normalized["ratings"] = []
        
        avg = normalized.get("average_score")
        if isinstance(avg, (int, float)):
            # Legacy values were stored on a 0-5 scale
            if avg > 1.5:
                normalized["average_score"] = avg / 5
        else:
            normalized["average_score"] = 0.0
        
        section_totals, total_score, total_possible = self.compute_section_summary(normalized.get("ratings", []))
        normalized["section_totals"] = section_totals
        normalized["total_score"] = total_score
        normalized["total_possible"] = total_possible
        if total_possible:
            normalized["average_score"] = total_score / total_possible
        normalized["score_percentage"] = normalized.get("average_score", 0.0) * 100
        normalized["rated_item_count"] = sum(int(entry.get("possible", 0) / 5) for entry in section_totals)
        
        counts = normalized.get("score_counts") or {}
        for label in self.rating_value_map.keys():
            counts.setdefault(label, 0)
        normalized["score_counts"] = counts
        
        recommendation = normalized.get("recommendation")
        if recommendation == self.recommendation_default:
            normalized["recommendation"] = ""
        
        return normalized
    
    def submit_evaluation(self):
        """Submit evaluation and save to system"""
        # Validate required fields
        if not self.entries["evaluator_name"].get():
            messagebox.showerror("Error", "Please enter Evaluator Name")
            self.notebook.select(0)
            return
        
        if not self.entries["trainer_name"].get():
            messagebox.showerror("Error", "Please enter Trainer Name")
            self.notebook.select(0)
            return
        
        if self.recommendation.get() == self.recommendation_default:
            messagebox.showerror("Error", "Please select an overall recommendation")
            self.notebook.select(0)
            return
        
        eval_data, summary = self.build_form_evaluation_data()
        total_score = summary.get("total_score", 0)
        total_possible = summary.get("total_possible", 0)
        average_percent = summary.get("average", 0)
        average_display = f"{average_percent:.0%}" if total_possible else "N/A"
        overall_raw = f"{total_score:.0f}/{total_possible}" if total_possible else "0/0"
        eval_data["submission_date"] = datetime.now().strftime('%m/%d/%Y %I:%M %p')
        
        # Create trainer directory
        trainer_dir = os.path.join(self.data_dir, 
                                  self.sanitize_filename(eval_data["trainer_name"]))
        if not os.path.exists(trainer_dir):
            os.makedirs(trainer_dir)
        
        # Save evaluation
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{self.sanitize_filename(eval_data['evaluator_name'])}_{timestamp}.json"
        filepath = os.path.join(trainer_dir, filename)
        
        with open(filepath, 'w') as f:
            json.dump(eval_data, f, indent=2)
        
        messagebox.showinfo(
            "Success",
            f"✅ Evaluation submitted successfully!\n\n"
            f"Evaluator: {eval_data['evaluator_name']}\n"
            f"Trainer: {eval_data['trainer_name']}\n"
            f"Average Score: {average_display} ({overall_raw})\n"
            f"Rated Items: {eval_data['rated_item_count']}\n\n"
            "This evaluation is now securely stored and can only be accessed by administrators."
        )
        
        # Ask if want to clear form
        if messagebox.askyesno("New Evaluation", "Would you like to start a new evaluation?"):
            self.clear_form()
    
    def export_report_word(self):
        """Export the current evaluation as a Word document."""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Unavailable", "The python-docx package is required to export Word reports.\n\nInstall with: pip install python-docx")
            return
        if not self.entries["evaluator_name"].get():
            messagebox.showerror("Error", "Please enter Evaluator Name before exporting.")
            return
        eval_data, summary = self.build_form_evaluation_data()
        context = self.prepare_report_context(eval_data, prepared_on=datetime.now())
        default_name = self.default_report_filename(eval_data.get("evaluator_name"), "docx")
        self.save_word_report(context, default_name)

    def export_report_pdf(self):
        """Export the current evaluation as a PDF document."""
        if not REPORTLAB_AVAILABLE:
            messagebox.showerror("Unavailable", "The reportlab package is required to export PDF reports.\n\nInstall with: pip install reportlab")
            return
        if not self.entries["evaluator_name"].get():
            messagebox.showerror("Error", "Please enter Evaluator Name before exporting.")
            return
        eval_data, summary = self.build_form_evaluation_data()
        context = self.prepare_report_context(eval_data, prepared_on=datetime.now())
        default_name = self.default_report_filename(eval_data.get("evaluator_name"), "pdf")
        self.save_pdf_report(context, default_name)

    def default_report_filename(self, evaluator_name, extension):
        """Generate a default filename for exports."""
        sanitized = self.sanitize_filename(evaluator_name or "Report") or "Report"
        timestamp = datetime.now().strftime('%Y%m%d')
        return f"HTASO_Evaluation_{sanitized}_{timestamp}.{extension}"

    def save_word_report(self, context, default_filename):
        """Prompt user to save a Word report."""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Unavailable", "The python-docx package is required to export Word reports.\n\nInstall with: pip install python-docx")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=default_filename
        )
        if not path:
            return
        try:
            self.write_word_report(context, path)
            messagebox.showinfo("Success", f"Word report exported successfully!\n\n{path}")
        except Exception as exc:
            messagebox.showerror("Export Failed", f"Unable to generate Word report.\n\n{exc}")

    def save_pdf_report(self, context, default_filename):
        """Prompt user to save a PDF report."""
        if not REPORTLAB_AVAILABLE:
            messagebox.showerror("Unavailable", "The reportlab package is required to export PDF reports.\n\nInstall with: pip install reportlab")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF Document", "*.pdf")],
            initialfile=default_filename
        )
        if not path:
            return
        try:
            self.write_pdf_report(context, path)
            messagebox.showinfo("Success", f"PDF report exported successfully!\n\n{path}")
        except Exception as exc:
            messagebox.showerror("Export Failed", f"Unable to generate PDF report.\n\n{exc}")

    def write_word_report(self, context, filepath):
        """Compose and save a formal Word report."""
        doc = Document()
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal'].font.size = Pt(11)

        if self.logo_path.exists():
            doc.add_picture(str(self.logo_path), width=Inches(1.5))
            doc.paragraphs[-1].alignment = 1
            doc.add_paragraph("")

        doc.add_heading("HTASO Umpire Evaluation Report", 0)
        doc.add_paragraph(f"Prepared on {context['prepared_on'].strftime('%B %d, %Y')}")
        doc.add_paragraph("")

        details = [
            ("Evaluator Name", context.get("evaluator_name", "N/A")),
            ("Trainer Name", context.get("trainer_name", "N/A")),
            ("Training Date", context.get("training_date", "N/A")),
            ("Observation Date", context.get("observation_date", "N/A")),
            ("Location", context.get("training_location", "N/A")),
            ("Submission Date", context.get("submission_date", "N/A")),
            ("Rated Items", str(context.get("rated_items", 0))),
            ("Average Score", f"{context.get('average_display', 'N/A')} ({context.get('overall_raw', '0/0')})")
        ]

        detail_table = doc.add_table(rows=0, cols=2)
        detail_table.style = "Table Grid"
        for label, value in details:
            row = detail_table.add_row().cells
            row[0].text = label
            row[1].text = value or "N/A"

        if context.get("section_totals"):
            doc.add_heading("Section Scores", level=1)
            score_table = doc.add_table(rows=1, cols=3)
            score_table.style = "Table Grid"
            header = score_table.rows[0].cells
            header[0].text = "Section"
            header[1].text = "Raw Score"
            header[2].text = "Percentage"
            for entry in context["section_totals"]:
                row = score_table.add_row().cells
                row[0].text = entry.get("section", "General")
                raw_text = f"{entry.get('score', 0):.0f}/{entry.get('possible', 0)}" if entry.get("possible") else "0/0"
                percent_text = f"{entry.get('percentage', 0):.0%}" if entry.get("possible") else "N/A"
                row[1].text = raw_text
                row[2].text = percent_text

        doc.add_heading("Evaluation Summary", level=1)
        doc.add_paragraph(
            f"Overall Score: {context.get('overall_raw', '0/0')} "
            f"({context.get('average_display', 'N/A')})."
        )

        for section, subsections in context.get("grouped_ratings", {}).items():
            doc.add_heading(section, level=1)
            for subsection, items in subsections.items():
                doc.add_heading(subsection, level=2)
                for item in items:
                    selection = item.get("selection") or "Not Rated"
                    prompt = item.get("prompt", "")
                    paragraph = doc.add_paragraph(style='List Bullet')
                    paragraph.add_run(f"{prompt} ({selection})")

        doc.add_heading("Evaluator Comments", level=1)
        for title, text in context.get("comments", {}).items():
            doc.add_heading(title, level=2)
            doc.add_paragraph(text if text else "None provided.")

        doc.save(filepath)

    def write_pdf_report(self, context, filepath):
        """Compose and save a formal PDF report."""
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=40,
            leftMargin=40,
            topMargin=60,
            bottomMargin=40
        )
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="Heading1Center", parent=styles['Heading1'], alignment=1))
        story = []

        if self.logo_path.exists():
            img = Image(str(self.logo_path), width=120, height=120)
            img.hAlign = 'CENTER'
            story.append(img)
            story.append(Spacer(1, 12))

        story.append(Paragraph("HTASO Umpire Evaluation Report", styles["Heading1Center"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Prepared on {context['prepared_on'].strftime('%B %d, %Y')}", styles["Normal"]))
        story.append(Spacer(1, 18))

        details_data = [
            ["Evaluator Name", context.get("evaluator_name", "N/A")],
            ["Trainer Name", context.get("trainer_name", "N/A")],
            ["Training Date", context.get("training_date", "N/A")],
            ["Observation Date", context.get("observation_date", "N/A")],
            ["Location", context.get("training_location", "N/A")],
            ["Submission Date", context.get("submission_date", "N/A")],
            ["Rated Items", str(context.get("rated_items", 0))],
            ["Average Score", f"{context.get('average_display', 'N/A')} ({context.get('overall_raw', '0/0')})"]
        ]
        details_table = Table(details_data, colWidths=[180, 360])
        details_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1D3557")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
        ]))
        story.append(details_table)
        story.append(Spacer(1, 18))

        if context.get("section_totals"):
            story.append(Paragraph("Section Scores", styles["Heading2"]))
            section_data = [["Section", "Raw Score", "Percentage"]]
            for entry in context["section_totals"]:
                raw_text = f"{entry.get('score', 0):.0f}/{entry.get('possible', 0)}" if entry.get("possible") else "0/0"
                percent_text = f"{entry.get('percentage', 0):.0%}" if entry.get("possible") else "N/A"
                section_data.append([
                    entry.get("section", "General"),
                    raw_text,
                    percent_text
                ])
            section_table = Table(section_data, colWidths=[200, 160, 180])
            section_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2A9D8F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("BOX", (0, 0), (-1, -1), 0.75, colors.grey),
            ]))
            story.append(section_table)
            story.append(Spacer(1, 18))

        story.append(Paragraph(
            f"Overall Score: {context.get('overall_raw', '0/0')} "
            f"({context.get('average_display', 'N/A')}).",
            styles["Normal"]
        ))
        story.append(Spacer(1, 18))

        for section, subsections in context.get("grouped_ratings", {}).items():
            story.append(Paragraph(section, styles["Heading2"]))
            for subsection, items in subsections.items():
                story.append(Paragraph(subsection, styles["Heading3"]))
                for item in items:
                    selection = item.get("selection") or "Not Rated"
                    prompt = item.get("prompt", "")
                    story.append(Paragraph(f"• {prompt} ({selection})", styles["Normal"]))
                story.append(Spacer(1, 6))
            story.append(Spacer(1, 12))

        story.append(Paragraph("Evaluator Comments", styles["Heading2"]))
        for title, text in context.get("comments", {}).items():
            story.append(Paragraph(title, styles["Heading3"]))
            cleaned = (text or "None provided.").replace("\n", "<br/>")
            story.append(Paragraph(cleaned, styles["Normal"]))
        story.append(Spacer(1, 6))

        doc.build(story)
    
    def export_report_word_from_record(self, record):
        """Export a stored evaluation record as a Word report."""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Unavailable", "The python-docx package is required to export Word reports.\n\nInstall with: pip install python-docx")
            return
        context = self.prepare_report_context(record, prepared_on=datetime.now())
        default_name = self.default_report_filename(context.get("evaluator_name"), "docx")
        self.save_word_report(context, default_name)

    def export_report_pdf_from_record(self, record):
        """Export a stored evaluation record as a PDF report."""
        if not REPORTLAB_AVAILABLE:
            messagebox.showerror("Unavailable", "The reportlab package is required to export PDF reports.\n\nInstall with: pip install reportlab")
            return
        context = self.prepare_report_context(record, prepared_on=datetime.now())
        default_name = self.default_report_filename(context.get("evaluator_name"), "pdf")
        self.save_pdf_report(context, default_name)
    
    def admin_login(self):
        """Show admin login dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Admin Access")
        dialog.geometry("400x250")
        dialog.configure(bg=self.colors['bg'])
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center dialog
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (dialog.winfo_width() // 2)
        y = (dialog.winfo_screenheight() // 2) - (dialog.winfo_height() // 2)
        dialog.geometry(f"+{x}+{y}")
        
        tk.Label(dialog, text="🔐 Administrator Login", font=("Segoe UI", 16, "bold"),
                bg=self.colors['bg'], fg=self.colors['text']).pack(pady=20)
        
        tk.Label(dialog, text="Enter Admin Password:", font=("Segoe UI", 10),
                bg=self.colors['bg'], fg=self.colors['text']).pack(pady=10)
        
        password_entry = tk.Entry(dialog, show="*", font=("Segoe UI", 12),
                                 width=25, relief="solid", bd=1)
        password_entry.pack(pady=10)
        password_entry.focus()
        
        def login():
            password = password_entry.get()
            if self.verify_admin_password(password):
                dialog.destroy()
                self.show_admin_panel()
            else:
                messagebox.showerror("Access Denied", "Invalid admin password")
                password_entry.delete(0, tk.END)
        
        password_entry.bind("<Return>", lambda e: login())
        
        btn_frame = tk.Frame(dialog, bg=self.colors['bg'])
        btn_frame.pack(pady=20)
        
        login_btn = tk.Button(btn_frame, text="Login", command=login,
                             font=("Segoe UI", 10, "bold"),
                             bg=self.colors['primary'], fg="white",
                             padx=30, pady=10, relief="flat", cursor="hand2")
        login_btn.pack(side="left", padx=5)
        
        tk.Button(btn_frame, text="Cancel", command=dialog.destroy,
                 font=("Segoe UI", 10), bg="#95A5A6", fg="white",
                 padx=30, pady=10, relief="flat", cursor="hand2").pack(side="left", padx=5)
    
    def show_admin_panel(self):
        """Show admin panel with all evaluations"""
        if self.admin_window and self.admin_window.winfo_exists():
            self.admin_window.deiconify()
            self.admin_window.lift()
            self.admin_window.focus_force()
            return
        
        admin_window = tk.Toplevel(self.root)
        self.admin_window = admin_window
        admin_window.title("Admin Panel - Evaluation Reports")
        admin_window.geometry("1200x700")
        admin_window.configure(bg=self.colors['bg'])
        admin_window.protocol("WM_DELETE_WINDOW", self._close_admin_window)
         
        # Header
        header = tk.Frame(admin_window, bg=self.colors['primary'], height=70)
        header.pack(fill="x")
        header.pack_propagate(False)
         
        tk.Label(header, text="Administrator Panel", font=("Segoe UI", 20, "bold"),
                fg="white", bg=self.colors['primary']).pack(side="left", padx=30, pady=15)
        
        # Main content
        content = tk.Frame(admin_window, bg=self.colors['bg'])
        content.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Trainer selection
        select_frame = tk.Frame(content, bg=self.colors['card'], relief="solid", bd=1)
        select_frame.pack(fill="x", pady=(0, 20))
        
        tk.Label(select_frame, text="Select Evaluator:", font=("Segoe UI", 11, "bold"),
                bg=self.colors['card'], fg=self.colors['text']).pack(side="left", padx=20, pady=15)
        
        all_evaluations = []
        if os.path.exists(self.data_dir):
            for trainer in os.listdir(self.data_dir):
                trainer_path = os.path.join(self.data_dir, trainer)
                if not os.path.isdir(trainer_path):
                    continue
                for entry in os.listdir(trainer_path):
                    if not entry.endswith(".json"):
                        continue
                    filepath = os.path.join(trainer_path, entry)
                    try:
                        with open(filepath, 'r') as f:
                            raw = json.load(f)
                        record = self.normalize_evaluation_data(raw)
                        all_evaluations.append({
                            "filepath": filepath,
                            "trainer": trainer,
                            "evaluator": record.get("evaluator_name", "Unknown"),
                            "submission": record.get("submission_date", "N/A"),
                            "submission_display": self.format_datetime_display(record.get("submission_date", "N/A")),
                            "score_display": f"{record.get('average_score', 0):.0%}",
                            "recommendation": record.get("recommendation") or "Not Selected"
                        })
                    except Exception:
                        continue
        
        if not all_evaluations:
            tk.Label(select_frame, text="No evaluations found", font=("Segoe UI", 10),
                    fg=self.colors['text_light'], bg=self.colors['card']).pack(side="left", padx=10)
        else:
            evaluator_names = sorted({item["evaluator"] for item in all_evaluations})
            evaluator_choices = ["All Evaluators"] + evaluator_names
            evaluator_var = tk.StringVar(value=evaluator_choices[0])
            evaluator_combo = ttk.Combobox(select_frame, textvariable=evaluator_var,
                                          values=evaluator_choices, state="readonly",
                                          font=("Segoe UI", 10), width=30)
            evaluator_combo.pack(side="left", padx=10)
            
            def populate_tree(selected):
                # Clear existing items
                for item in tree.get_children():
                    tree.delete(item)
                for record in all_evaluations:
                    if selected != "All Evaluators" and record["evaluator"] != selected:
                        continue
                    tree.insert(
                        "",
                        "end",
                        values=(
                            record["trainer"],
                            record["evaluator"],
                            record.get("submission_display", record["submission"]),
                            record["score_display"],
                            record["recommendation"],
                            record["filepath"]
                        )
                    )
            
            def load_evaluations():
                populate_tree(evaluator_var.get())
            
            evaluator_combo.bind("<<ComboboxSelected>>", lambda e: populate_tree(evaluator_var.get()))
            
            tk.Button(select_frame, text="Load", command=load_evaluations,
                     font=("Segoe UI", 10), bg=self.colors['primary'], fg="white",
                     relief="flat", padx=20, pady=8, cursor="hand2").pack(side="left", padx=10)
            initial_selection = evaluator_var.get()
        
        # Evaluations list
        list_frame = tk.Frame(content, bg=self.colors['card'], relief="solid", bd=1)
        list_frame.pack(fill="both", expand=True)
        
        tk.Label(list_frame, text="Submitted Evaluations", font=("Segoe UI", 12, "bold"),
                bg=self.colors['card'], fg=self.colors['text']).pack(anchor="w", padx=20, pady=10)
        
        # Treeview
        tree_frame = tk.Frame(list_frame, bg=self.colors['card'])
        tree_frame.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        columns = ('Trainer', 'Evaluator', 'Date', 'Score', 'Recommendation')
        tree = ttk.Treeview(tree_frame, columns=columns, show='headings', height=15)
        
        tree.heading('Trainer', text='Trainer Name')
        tree.heading('Evaluator', text='Evaluator Name')
        tree.heading('Date', text='Submission Date')
        tree.heading('Score', text='Avg Score')
        tree.heading('Recommendation', text='Recommendation')
        
        tree.column('Trainer', width=220)
        tree.column('Evaluator', width=220)
        tree.column('Date', width=160)
        tree.column('Score', width=100, anchor="center")
        tree.column('Recommendation', width=320)
        
        scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)
        
        tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        self.enable_mousewheel(tree)
        tree.bind("<MouseWheel>", lambda e: tree.yview_scroll(int(-1 * (e.delta / 120)), "units"), add="+")
        tree.bind("<Button-4>", lambda e: tree.yview_scroll(-1, "units"), add="+")
        tree.bind("<Button-5>", lambda e: tree.yview_scroll(1, "units"), add="+")
        if all_evaluations:
            populate_tree(initial_selection)
        
        # Action buttons
        action_frame = tk.Frame(list_frame, bg=self.colors['card'])
        action_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        def view_evaluation():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select an evaluation to view")
                return
            
            item = tree.item(selection[0])
            filepath = item['values'][-1]
            
            with open(filepath, 'r') as f:
                data = json.load(f)
            
            self.show_evaluation_detail(data, admin_window)
        
        def get_selected_record():
            selection = tree.selection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select an evaluation first.")
                return None
            
            item = tree.item(selection[0])
            filepath = item['values'][-1]
            
            with open(filepath, 'r') as f:
                data = json.load(f)
            
            return data
        
        def export_word_selected():
            record = get_selected_record()
            if record is None:
                return
            self.export_report_word_from_record(record)
        
        def export_pdf_selected():
            record = get_selected_record()
            if record is None:
                return
            self.export_report_pdf_from_record(record)
        
        tk.Button(action_frame, text="👁️ View Details", command=view_evaluation,
                 font=("Segoe UI", 10), bg=self.colors['primary'], fg="white",
                 relief="flat", padx=20, pady=10, cursor="hand2").pack(side="left", padx=5)
        
        tk.Button(action_frame, text="📄 Export Word", command=export_word_selected,
                 font=("Segoe UI", 10), bg=self.colors['secondary'], fg="white",
                 relief="flat", padx=20, pady=10, cursor="hand2").pack(side="left", padx=5)
        
        tk.Button(action_frame, text="📃 Export PDF", command=export_pdf_selected,
                 font=("Segoe UI", 10), bg=self.colors['accent'], fg="white",
                 relief="flat", padx=20, pady=10, cursor="hand2").pack(side="left", padx=5)
        
        tk.Button(action_frame, text="🔒 Change Password", 
                 command=lambda: self.change_admin_password(admin_window),
                 font=("Segoe UI", 10), bg="#E67E22", fg="white",
                 relief="flat", padx=20, pady=10, cursor="hand2").pack(side="left", padx=5)
        
        tk.Button(action_frame, text="❌ Close", command=self._close_admin_window,
                 font=("Segoe UI", 10), bg="#95A5A6", fg="white",
                 relief="flat", padx=20, pady=10, cursor="hand2").pack(side="right", padx=5)
    
    def _close_admin_window(self):
        """Safely close the admin window."""
        if self.admin_window and self.admin_window.winfo_exists():
            self.admin_window.destroy()
        self.admin_window = None
    
    def show_evaluation_detail(self, data, parent):
        """Show detailed view of an evaluation"""
        detail_window = tk.Toplevel(parent)
        detail_window.title(f"Evaluation Details - {data['evaluator_name']}")
        detail_window.geometry("900x700")
        detail_window.configure(bg=self.colors['bg'])

        record = self.normalize_evaluation_data(data)

        canvas = tk.Canvas(detail_window, bg=self.colors['bg'], highlightthickness=0)
        scrollbar = ttk.Scrollbar(detail_window, orient="vertical", command=canvas.yview)
        frame = tk.Frame(canvas, bg=self.colors['bg'])

        frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        self.enable_mousewheel(canvas)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        section_summary = record.get("section_totals") or []
        total_score = record.get("total_score")
        total_possible = record.get("total_possible")
        if not section_summary or total_score is None or total_possible is None:
            section_summary, total_score, total_possible = self.compute_section_summary(record.get("ratings", []))
        average_score = (total_score / total_possible) if total_possible else 0
        overall_raw = f"{total_score:.0f}/{total_possible}" if total_possible else "0/0"
        average_display = f"{average_score:.0%}" if total_possible else "N/A"
        rated_count = record.get("rated_item_count", 0)

        info_card = self.create_card(frame, "Basic Information")
        info_card.pack(fill="x", padx=20, pady=10)
        info_lines = [
            f"Evaluator: {record['evaluator_name']}",
            f"Trainer: {record['trainer_name']}",
            f"Training Date: {record.get('training_date', 'N/A')}",
            f"Location: {record.get('training_location', 'N/A')}",
            f"Submission Date: {self.format_datetime_display(record.get('submission_date', 'N/A'))}",
            f"Rated Items: {rated_count}",
            f"Average Score: {average_display} ({overall_raw})"
        ]
        tk.Label(
            info_card.body,
            text="\n".join(info_lines),
            font=("Segoe UI", 10),
            justify="left",
            bg=self.colors['card'],
            fg=self.colors['text']
        ).pack(fill="x")

        if section_summary:
            section_card = self.create_card(
                frame,
                "Section Scores",
                subtitle="Raw totals with percentage for each evaluation group."
            )
            section_card.pack(fill="x", padx=20, pady=10)
            for entry in section_summary:
                possible = entry.get("possible", 0)
                score_value = entry.get("score", 0)
                raw_text = f"{score_value:.0f}/{possible}" if possible else "0/0"
                percent_text = f"{entry.get('percentage', 0):.0%}" if possible else "N/A"
                tk.Label(
                    section_card.body,
                    text=f"{entry.get('section', 'General')}: {raw_text} ({percent_text})",
                    font=("Segoe UI", 10),
                    bg=self.colors['card'],
                    fg=self.colors['text']
                ).pack(anchor="w", pady=2)

        ratings = record.get("ratings", [])
        ratings_card = self.create_card(frame, "Ratings Summary")
        ratings_card.pack(fill="x", padx=20, pady=10)

        grouped = OrderedDict()
        for item in ratings:
            section_name = item.get("section") or "General"
            subsection_name = item.get("subsection") or "Criteria"
            grouped.setdefault(section_name, OrderedDict()).setdefault(subsection_name, []).append(item)

        for section, subsections in grouped.items():
            tk.Label(
                ratings_card.body,
                text=section,
                font=("Segoe UI Semibold", 12),
                bg=self.colors['card'],
                fg=self.colors['text']
            ).pack(anchor="w", pady=(8, 4))
            for subsection, items in subsections.items():
                tk.Label(
                    ratings_card.body,
                    text=f"  {subsection}",
                    font=("Segoe UI", 10, "bold"),
                    bg=self.colors['card'],
                    fg=self.colors['text_light']
                ).pack(anchor="w", pady=(2, 2))
                for item in items:
                    selection = item.get("selection") or "Not Rated"
                    display = f"     - {item.get('prompt', '')} ({selection})"
                    tk.Label(
                        ratings_card.body,
                        text=display,
                        font=("Segoe UI", 10),
                        bg=self.colors['card'],
                        fg=self.colors['text'],
                        wraplength=820,
                        justify="left"
                    ).pack(anchor="w", pady=1)

        rec_card = self.create_card(frame, "Recommendation")
        rec_card.pack(fill="x", padx=20, pady=10)
        tk.Label(
            rec_card.body,
            text=record.get('recommendation') or "Not Selected",
            font=("Segoe UI", 11, "bold"),
            bg=self.colors['card'],
            fg=self.colors['secondary'],
            wraplength=820,
            justify="left"
        ).pack(fill="x")

        comments_map = [
            ("Strengths", "strengths"),
            ("Areas for Improvement", "improvement"),
            ("Development Recommendations", "development"),
            ("Overall Comments", "overall")
        ]
        for title, key in comments_map:
            comment_text = record.get('comments', {}).get(key)
            if comment_text:
                comment_card = self.create_card(frame, title)
                comment_card.pack(fill="x", padx=20, pady=10)
                tk.Label(
                    comment_card.body,
                    text=comment_text,
                    font=("Segoe UI", 10),
                    bg=self.colors['card'],
                    fg=self.colors['text'],
                    wraplength=820,
                    justify="left"
                ).pack(anchor="w")

    def change_admin_password(self, parent):
        """Change admin password"""
        dialog = tk.Toplevel(parent)
        dialog.title("Change Admin Password")
        dialog.geometry("400x300")
        dialog.configure(bg=self.colors['bg'])
        dialog.transient(parent)
        dialog.grab_set()
        
        tk.Label(dialog, text="Change Password", font=("Segoe UI", 14, "bold"),
                bg=self.colors['bg'], fg=self.colors['text']).pack(pady=20)
        
        tk.Label(dialog, text="Current Password:", font=("Segoe UI", 10),
                bg=self.colors['bg']).pack()
        current_pw = tk.Entry(dialog, show="*", font=("Segoe UI", 11), width=25)
        current_pw.pack(pady=5)
        
        tk.Label(dialog, text="New Password:", font=("Segoe UI", 10),
                bg=self.colors['bg']).pack()
        new_pw = tk.Entry(dialog, show="*", font=("Segoe UI", 11), width=25)
        new_pw.pack(pady=5)
        
        tk.Label(dialog, text="Confirm New Password:", font=("Segoe UI", 10),
                bg=self.colors['bg']).pack()
        confirm_pw = tk.Entry(dialog, show="*", font=("Segoe UI", 11), width=25)
        confirm_pw.pack(pady=5)
        
        def change():
            if not self.verify_admin_password(current_pw.get()):
                messagebox.showerror("Error", "Current password is incorrect")
                return
            
            if new_pw.get() != confirm_pw.get():
                messagebox.showerror("Error", "New passwords do not match")
                return
            
            if len(new_pw.get()) < 6:
                messagebox.showerror("Error", "Password must be at least 6 characters")
                return
            
            new_hash = hashlib.sha256(new_pw.get().encode()).hexdigest()
            with open(self.admin_file, 'w') as f:
                json.dump({"password_hash": new_hash}, f)
            
            messagebox.showinfo("Success", "Password changed successfully!")
            dialog.destroy()
        
        tk.Button(dialog, text="Change Password", command=change,
                 font=("Segoe UI", 10), bg=self.colors['primary'], fg="white",
                 padx=20, pady=10, relief="flat", cursor="hand2").pack(pady=20)
    
    def sanitize_filename(self, filename):
        """Remove invalid characters from filename"""
        return "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).strip()
    
    def format_datetime_display(self, value, include_time=True):
        """Convert a stored datetime string into MM/DD/YYYY (optionally with time)."""
        if not value:
            return "N/A"
        patterns = [
            "%m/%d/%Y %I:%M %p",
            "%Y-%m-%d %H:%M:%S",
            "%Y-%m-%d %H:%M",
            "%Y/%m/%d %H:%M:%S",
            "%Y/%m/%d"
        ]
        for pattern in patterns:
            try:
                dt = datetime.strptime(value, pattern)
                return dt.strftime("%m/%d/%Y %I:%M %p") if include_time else dt.strftime("%m/%d/%Y")
            except ValueError:
                continue
        return value
    
    def open_date_picker(self, target_entry):
        """Show a simple calendar dialog and populate the target entry."""
        if not TKCALENDAR_AVAILABLE:
            messagebox.showinfo(
                "Calendar Unavailable",
                "Install the 'tkcalendar' package to enable the date picker.\n\npip install tkcalendar"
            )
            return
        
        try:
            current_value = target_entry.get()
            base_date = datetime.strptime(current_value, "%m/%d/%Y")
        except ValueError:
            base_date = datetime.today()
        
        picker = tk.Toplevel(self.root)
        picker.title("Select Date")
        picker.grab_set()
        picker.transient(self.root)
        
        cal = Calendar(
            picker,
            selectmode="day",
            year=base_date.year,
            month=base_date.month,
            day=base_date.day,
            date_pattern="mm/dd/yyyy"
        )
        cal.pack(padx=10, pady=10)
        
        def apply_selection():
            selected = cal.selection_get()
            if selected:
                target_entry.delete(0, tk.END)
                target_entry.insert(0, selected.strftime("%m/%d/%Y"))
            picker.destroy()
        
        tk.Button(
            picker,
            text="Select",
            command=apply_selection,
            font=("Segoe UI", 10),
            bg=self.colors['primary'],
            fg="white",
            relief="flat",
            padx=12,
            pady=6,
            cursor="hand2"
        ).pack(side="left", padx=10, pady=(0, 12))
        
        tk.Button(
            picker,
            text="Cancel",
            command=picker.destroy,
            font=("Segoe UI", 10),
            bg="#95A5A6",
            fg="white",
            relief="flat",
            padx=12,
            pady=6,
            cursor="hand2"
        ).pack(side="right", padx=10, pady=(0, 12))
        
        picker.focus_force()
    
    def create_evaluation_sheet(self, ws):
        """Create evaluation sheet for Excel export"""
        ws.column_dimensions['A'].width = 70
        ws.column_dimensions['B'].width = 25

        row = 1
        ws[f'A{row}'] = "EVALUATION & OBSERVATION RESULTS"
        ws.merge_cells(f'A{row}:B{row}')
        ws[f'A{row}'].font = Font(size=14, bold=True)
        ws[f'A{row}'].alignment = Alignment(horizontal='center')
        row += 2

        # Basic info
        for label, key in [("Evaluator Name:", "evaluator_name"),
                          ("Training Date:", "training_date"),
                          ("Trainer Name:", "trainer_name"),
                          ("Location:", "training_location")]:
            ws[f'A{row}'] = label
            ws[f'B{row}'] = self.entries[key].get()
            ws[f'A{row}'].font = Font(bold=True)
            row += 1

        summary = self.collect_ratings()

        row += 1
        ws[f'A{row}'] = "RATINGS"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        row += 1

        ws[f'A{row}'] = "Section / Criterion"
        ws[f'B{row}'] = "Selection"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].font = Font(bold=True)
        row += 1

        for item in summary["entries"]:
            section_label = item.get('section') or "General"
            subsection_label = item.get('subsection')
            descriptor = section_label
            if subsection_label:
                descriptor += f" -> {subsection_label}"
            prompt = item.get('prompt', '')
            if prompt:
                descriptor += f" - {prompt}"
            ws[f'A{row}'] = descriptor
            ws[f'A{row}'].alignment = Alignment(wrap_text=True)
            ws[f'B{row}'] = item.get('selection') or "Not Rated"
            row += 1

        row += 1
        ws[f'A{row}'] = "Average Score"
        ws[f'B{row}'] = f"{average_display} ({overall_raw})"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].font = Font(bold=True)
        row += 1
        ws[f'A{row}'] = "Rated Items"
        ws[f'B{row}'] = summary['rated_count']

        row += 2
        ws[f'A{row}'] = "SECTION SCORES"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        for entry in summary['section_totals']:
            possible = entry.get('possible', 0)
            score_value = entry.get('score', 0)
            raw_text = f"{score_value:.0f}/{possible}" if possible else "0/0"
            percent_text = f"{entry.get('percentage', 0):.0%}" if possible else "N/A"
            ws[f'A{row}'] = entry.get('section', 'General')
            ws[f'B{row}'] = f"{raw_text} ({percent_text})"
            row += 1

        row += 1
        ws[f'A{row}'] = "Overall Score"
        if total_possible:
            ws[f'B{row}'] = f"{total_score:.0f}/{total_possible} ({average_percent:.0%})"
        else:
            ws[f'B{row}'] = "N/A"

        row += 2
        ws[f'A{row}'] = "RECOMMENDATION"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        recommendation_display = self.recommendation.get()
        if recommendation_display == self.recommendation_default:
            recommendation_display = "Not Selected"
        ws[f'A{row}'] = recommendation_display
        ws.merge_cells(f'A{row}:B{row}')

    def create_comments_sheet(self, ws):
        """Create comments sheet for Excel export"""
        ws.column_dimensions['A'].width = 100
        
        row = 1
        ws[f'A{row}'] = "COMMENTS"
        ws[f'A{row}'].font = Font(size=14, bold=True)
        row += 2
        
        for title, text_widget in [("STRENGTHS:", self.strengths_text),
                                  ("IMPROVEMENTS:", self.improvement_text),
                                  ("DEVELOPMENT:", self.development_text),
                                  ("OVERALL:", self.overall_text)]:
            ws[f'A{row}'] = title
            ws[f'A{row}'].font = Font(bold=True, size=11)
            row += 1
            ws[f'A{row}'] = text_widget.get("1.0", "end-1c")
            ws[f'A{row}'].alignment = Alignment(wrap_text=True)
            row += 3
    
    def create_scale_sheet_excel(self, ws):
        """Create scale criteria sheet for Excel"""
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 80
        
        row = 1
        ws[f'A{row}'] = "RATING SCALE"
        ws.merge_cells(f'A{row}:B{row}')
        ws[f'A{row}'].font = Font(size=14, bold=True)
        row += 2
        
        scales = [
            ("5 - EXCELLENT", "Consistently exceeds expectations"),
            ("4 - VERY GOOD", "Meets and often exceeds expectations"),
            ("3 - SATISFACTORY", "Meets basic expectations"),
            ("2 - NEEDS IMPROVEMENT", "Partially meets expectations"),
            ("1 - UNSATISFACTORY", "Does not meet expectations")
        ]
        
        for rating, desc in scales:
            ws[f'A{row}'] = rating
            ws[f'B{row}'] = desc
            ws[f'A{row}'].font = Font(bold=True)
            row += 1
    
    def create_instructions_sheet_excel(self, ws):
        """Create instructions sheet for Excel"""
        ws.column_dimensions['A'].width = 100
        
        row = 1
        ws[f'A{row}'] = "INSTRUCTIONS"
        ws[f'A{row}'].font = Font(size=14, bold=True)
        row += 2
        
        ws[f'A{row}'] = "Complete all sections of the evaluation form based on observed performance."
        row += 1
        ws[f'A{row}'] = "Use the rating scale consistently and provide detailed comments."
        row += 1
    
    def clear_form(self):
        """Clear all form fields"""
        if messagebox.askyesno("Clear Form", "Are you sure you want to clear all fields?"):
            for entry in self.entries.values():
                entry.delete(0, tk.END)
            
            for var in self.ratings.values():
                var.set("Select result")
            
            self.strengths_text.delete("1.0", tk.END)
            self.improvement_text.delete("1.0", tk.END)
            self.development_text.delete("1.0", tk.END)
            self.overall_text.delete("1.0", tk.END)
            self.recommendation.set(self.recommendation_default)
            self.update_recommendation_indicators()
            
            self.notebook.select(0)

if __name__ == "__main__":
    root = tk.Tk()
    app = EvaluatorTrainingApp(root)
    root.mainloop()
