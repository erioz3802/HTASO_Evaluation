"""Word document export functionality using python-docx."""
from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches


def generate_word_report(data: Dict[str, Any], logo_path: Optional[Path] = None) -> BytesIO:
    """Generate a Word document report from evaluation data.

    Args:
        data: Evaluation data dictionary
        logo_path: Optional path to logo image

    Returns:
        BytesIO object containing the Word document
    """
    doc = Document()

    # Set default font
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)

    # Add logo if available
    if logo_path and logo_path.exists():
        try:
            doc.add_picture(str(logo_path), width=Inches(1.5))
            doc.paragraphs[-1].alignment = 1  # Center alignment
            doc.add_paragraph("")
        except Exception:
            pass  # Skip logo if there's an error

    # Title
    doc.add_heading("HTASO Umpire Evaluation Report", 0)

    # Prepared date
    prepared_on = datetime.now().strftime('%B %d, %Y')
    doc.add_paragraph(f"Prepared on {prepared_on}")
    doc.add_paragraph("")

    # Basic details table
    details = [
        ("Evaluator Name", data.get("evaluator_name", "N/A")),
        ("Trainer Name", data.get("trainer_name", "N/A")),
        ("Training Date", data.get("training_date", "N/A")),
        ("Observation Date", data.get("observation_date", "N/A") or "N/A"),
        ("Location", data.get("training_location", "N/A") or "N/A"),
        ("Type of Evaluation", data.get("eval_type", "N/A") or "N/A"),
        ("Submission Date", data.get("submission_date", "N/A")),
        ("Rated Items", str(data.get("rated_item_count", 0))),
    ]

    # Calculate HTASO average and rank
    total_score = data.get("total_score", 0)
    total_possible = data.get("total_possible", 0)
    if total_possible and total_possible > 0:
        htaso_avg = 6.0 - (total_score / (total_possible / 5.0))
        rank = round(htaso_avg)
        average_score = total_score / total_possible
        details.append(("HTASO Average", f"{htaso_avg:.2f}"))
        details.append(("Rank", str(rank)))
        details.append(("Percentage", f"{average_score:.0%}"))
    else:
        htaso_avg = None
        rank = None
        details.append(("HTASO Average", "N/A"))

    detail_table = doc.add_table(rows=0, cols=2)
    detail_table.style = "Table Grid"
    for label, value in details:
        row = detail_table.add_row().cells
        row[0].text = label
        row[1].text = value or "N/A"

    doc.add_paragraph("")

    # Section scores table
    section_totals = data.get("section_totals", [])
    if section_totals:
        doc.add_heading("Section Scores", level=1)
        score_table = doc.add_table(rows=1, cols=4)
        score_table.style = "Table Grid"
        header = score_table.rows[0].cells
        header[0].text = "Section"
        header[1].text = "HTASO Avg"
        header[2].text = "Rank"
        header[3].text = "Percentage"

        for entry in section_totals:
            row = score_table.add_row().cells
            row[0].text = entry.get("section", "General")
            sec_possible = entry.get("possible", 0)
            sec_score = entry.get("score", 0)
            if sec_possible and sec_possible > 0:
                sec_htaso_avg = 6.0 - (sec_score / (sec_possible / 5.0))
                sec_rank = round(sec_htaso_avg)
                row[1].text = f"{sec_htaso_avg:.2f}"
                row[2].text = str(sec_rank)
                row[3].text = f"{entry.get('percentage', 0):.0%}"
            else:
                row[1].text = "—"
                row[2].text = "—"
                row[3].text = "N/A"

        doc.add_paragraph("")

    # Evaluation summary
    doc.add_heading("Evaluation Summary", level=1)
    if htaso_avg is not None:
        doc.add_paragraph(f"Overall HTASO Average: {htaso_avg:.2f} | Rank: {rank}")
    else:
        doc.add_paragraph("Overall Score: N/A")

    # Ratings grouped by section
    ratings = data.get("ratings", [])
    if ratings:
        # Group ratings by section and subsection
        from collections import OrderedDict
        grouped = OrderedDict()
        for item in ratings:
            section_name = item.get("section") or "General"
            subsection_name = item.get("subsection") or "Criteria"
            if section_name not in grouped:
                grouped[section_name] = OrderedDict()
            if subsection_name not in grouped[section_name]:
                grouped[section_name][subsection_name] = []
            grouped[section_name][subsection_name].append(item)

        # Add ratings to document
        for section, subsections in grouped.items():
            doc.add_heading(section, level=1)
            for subsection, items in subsections.items():
                doc.add_heading(subsection, level=2)
                for item in items:
                    selection = item.get("selection") or "Not Rated"
                    prompt = item.get("prompt", "")
                    paragraph = doc.add_paragraph(style='List Bullet')
                    paragraph.add_run(f"{prompt} ({selection})")

    # Recommendation
    doc.add_heading("Overall Recommendation", level=1)
    recommendation = data.get("recommendation") or "Not Selected"
    doc.add_paragraph(recommendation)

    # Comments
    comments = data.get("comments", {})
    if comments:
        doc.add_heading("Evaluator Comments", level=1)
        comment_sections = [
            ("Strengths Observed", "strengths"),
            ("Areas for Improvement", "improvement"),
            ("Development Recommendations", "development"),
            ("Overall Assessment Comments", "overall")
        ]

        for title, key in comment_sections:
            text = comments.get(key, "")
            doc.add_heading(title, level=2)
            doc.add_paragraph(text if text else "None provided.")

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def get_word_filename(evaluator_name: str) -> str:
    """Generate a standard filename for Word export.

    Args:
        evaluator_name: Name of the evaluator

    Returns:
        Filename string
    """
    # Sanitize filename
    safe_name = "".join(c for c in evaluator_name if c.isalnum() or c in (' ', '-', '_')).strip()
    safe_name = safe_name or "Evaluation"
    timestamp = datetime.now().strftime('%Y%m%d')
    return f"HTASO_Evaluation_{safe_name}_{timestamp}.docx"
